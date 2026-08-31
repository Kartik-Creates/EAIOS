import io
import logging
from pathlib import Path

logger = logging.getLogger("eaios.document_parser")


class DocumentParserError(ValueError):
    """Raised when text extraction from a file fails or produces no usable content."""


def extract_text_from_file(filename: str, content_bytes: bytes) -> str:
    """Extract plain text from uploaded file bytes based on file extension.

    Supports:
      - .docx (Microsoft Word) via python-docx
      - .pdf (Adobe Portable Document Format) via pypdf
      - Plain text formats: .txt, .csv, .md, .json, .xml, .tsv, .log
    """
    if not content_bytes:
        raise DocumentParserError(f"File '{filename}' is empty.")

    ext = Path(filename).suffix.lower()

    if ext == ".docx":
        return _extract_from_docx(filename, content_bytes)
    elif ext == ".pdf":
        return _extract_from_pdf(filename, content_bytes)
    elif ext in {".txt", ".csv", ".tsv", ".md", ".json", ".xml", ".log", ".yaml", ".yml"}:
        return _extract_from_plaintext(filename, content_bytes)
    elif ext == ".doc":
        # python-docx does not support legacy binary DOC format (Word 97-2003).
        # We attempt basic text recovery or instruct the user to save as .docx or .pdf.
        try:
            return _extract_from_docx(filename, content_bytes)
        except Exception:
            text = _extract_from_plaintext(filename, content_bytes)
            # Check if there is discernible readable text
            words = [w for w in text.split() if w.isascii() and len(w) > 2]
            if len(words) >= 10:
                return text
            raise DocumentParserError(
                f"Legacy '.doc' format is not supported for '{filename}'. "
                "Please save and upload the document as modern '.docx' or '.pdf'."
            )
    else:
        # Fallback: attempt to decode as plaintext UTF-8
        try:
            text = _extract_from_plaintext(filename, content_bytes)
            if text.strip():
                return text
        except Exception:
            pass

        raise DocumentParserError(
            f"Unsupported file format '{ext}' for file '{filename}'. "
            "Supported formats: .docx, .pdf, .txt, .csv, .md, .json."
        )


def _extract_from_docx(filename: str, content_bytes: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise DocumentParserError("python-docx is not installed in this environment.") from exc

    try:
        doc = docx.Document(io.BytesIO(content_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract table text if present
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        all_text = "\n\n".join(paragraphs + table_texts).strip()
        if not all_text:
            raise DocumentParserError(f"No text content found in Word document '{filename}'.")
        return all_text
    except DocumentParserError:
        raise
    except Exception as exc:
        raise DocumentParserError(f"Failed to extract text from DOCX file '{filename}': {exc}") from exc


def _extract_from_pdf(filename: str, content_bytes: bytes) -> str:
    try:
        import pypdf
    except ImportError as exc:
        raise DocumentParserError("pypdf is not installed in this environment.") from exc

    try:
        reader = pypdf.PdfReader(io.BytesIO(content_bytes))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise DocumentParserError(f"PDF '{filename}' is password-protected and cannot be read.")

        extracted_pages = []
        for page_idx, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                extracted_pages.append(text.strip())

        full_text = "\n\n".join(extracted_pages).strip()
        if not full_text:
            raise DocumentParserError(f"No readable text found in PDF '{filename}' (it may be scanned/image-only).")
        return full_text
    except DocumentParserError:
        raise
    except Exception as exc:
        raise DocumentParserError(f"Failed to extract text from PDF file '{filename}': {exc}") from exc


def _extract_from_plaintext(filename: str, content_bytes: bytes) -> str:
    # Try utf-8 first, fallback to latin-1
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = content_bytes.decode(enc)
            if text.strip():
                return text.strip()
        except UnicodeDecodeError:
            continue

    raise DocumentParserError(f"Failed to decode text content from '{filename}'.")
